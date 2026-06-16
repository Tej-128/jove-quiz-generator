"""
JoVE Quiz Generator - Core Engine v3.1
Standalone quiz generation engine for 3 sets x 40 questions from PTx and Transcript docx files.

Primary changes in v3.1:
- Strict filename detection: *_PTx_NS.docx and *_Transcript.docx only.
- No hardcoded chapter-name fallback. Chapter name must come from PTx content.
- Validator rejects malformed or over-quota questions instead of exporting them.
- Per-set cap of exactly 40 accepted questions when generation succeeds.
- Top-up calls request only missing types and accept only requested types.
- Duplicate prevention uses full normalized question content, options, and near-match checks.
- Lesson IDs are validated against uploaded filename prefixes.
- Summary report includes skipped files, empty lessons, dropped/fallback types, and warnings.
- v3.1 tightens dropdown numbering, dedup cleanup after fallback, distribution warnings, and read-error reporting.
"""

from __future__ import annotations

import difflib
import hashlib
import json
import re
from pathlib import Path
from typing import Any, Callable

from docx import Document

# -----------------------------------------------------------------------------
# Constants
# -----------------------------------------------------------------------------

QUESTION_TYPES = [
    "Single Correct",
    "Multi Correct",
    "True or False",
    "Fill in the Blanks",
    "Dropdown",
    "Match the following",
    "Categorisation",
]

QUESTIONS_PER_SET = 40
NUM_SETS = 3
MIN_PER_TYPE = 5
MAX_RETRIES = 3

# Order required by the product spec. Single Correct is intentionally absent.
FALLBACK_PRIORITY = [
    "Match the following",
    "Categorisation",
    "Dropdown",
    "Fill in the Blanks",
    "True or False",
    "Multi Correct",
]

# Snippets are only for LLM prompt exclusion. Real dedup uses full signatures.
EXCLUSION_SNIPPET_CHARS = 280
NEAR_DUPLICATE_THRESHOLD = 0.90
SHORT_NEAR_DUPLICATE_THRESHOLD = 0.92
SHORT_NEAR_DUPLICATE_MIN_CHARS = 8
GROUNDING_MIN_OVERLAP = 0.25

REQUIRED_FIELDS = [
    "lesson_id",
    "question_index",
    "question_content",
    "question_type",
    "option_1",
    "option_2",
    "option_3",
    "option_4",
    "right_answer",
]

STOPWORDS = {
    "about", "above", "after", "again", "against", "along", "also", "among",
    "because", "before", "being", "between", "blank", "cannot", "choose", "correct",
    "could", "describe", "during", "each", "following", "from", "given", "have",
    "into", "more", "most", "only", "option", "question", "should", "statement",
    "than", "that", "their", "then", "there", "these", "this", "those", "through",
    "true", "false", "using", "what", "when", "where", "which", "while", "with",
    "would", "your", "dropdown", "match", "categorise", "categorize",
}

# -----------------------------------------------------------------------------
# File reading and lesson parsing
# -----------------------------------------------------------------------------


def read_docx(path: str, return_error: bool = False):
    """
    Extract text from paragraphs and tables.

    By default this preserves the old public API and returns only text. When
    return_error=True, returns (text, error_message) so callers can distinguish
    a legitimately empty DOCX from a read failure.
    """
    try:
        doc = Document(path)
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    # Use tabs, not pipes. Pipes are reserved quiz syntax downstream.
                    parts.append("\t".join(cells))
        text = "\n".join(parts)
        return (text, "") if return_error else text
    except Exception as exc:
        error = f"{type(exc).__name__}: {exc}"
        return ("", error) if return_error else ""


def _filename_lesson_id(stem: str) -> str:
    first = stem.split("_", 1)[0]
    return first if first.isdigit() else ""


def _is_ptx_filename(stem: str) -> bool:
    return stem.lower().endswith("_ptx_ns")


def _is_transcript_filename(stem: str) -> bool:
    return stem.lower().endswith("_transcript")


def _extract_chapter_name_from_pt(pt_text: str) -> str:
    """
    Extract the chapter-name value from PTx text without cosmetic rewriting.
    Supported labels include: Chapter Name:, Chapter Title:, Chapter:.
    If no explicit label exists, return an empty string rather than guessing.
    """
    lines = [line.strip() for line in pt_text.splitlines() if line.strip()]
    label_pattern = re.compile(
        r"^chapter\s*(?:name|title)?\s*[:\-]\s*(.+)$",
        flags=re.IGNORECASE,
    )

    for line in lines[:15]:
        match = label_pattern.match(line)
        if match:
            return match.group(1).strip()

    standalone_label = re.compile(r"^chapter\s*(?:name|title)?$", flags=re.IGNORECASE)
    for idx, line in enumerate(lines[:15]):
        if standalone_label.match(line) and idx + 1 < len(lines):
            return lines[idx + 1].strip()

    return ""


def parse_lesson_files(uploaded_files: list[dict[str, str]], return_report: bool = False):
    """
    Parse uploaded docx records into lesson records.

    Args:
        uploaded_files: list of {name: str, path: str} records.
        return_report: when True, returns (lessons, parse_report). When False, returns lessons.

    Lesson record keys:
        lesson_id, pt_text, transcript_text, chapter_name, pt_filename, transcript_filename
    """
    lessons: dict[str, dict[str, Any]] = {}
    report: dict[str, Any] = {
        "warnings": [],
        "skipped_files": [],
        "chapter_names": [],
        "chapter_name": "",
    }

    for file_record in uploaded_files:
        name = file_record.get("name", "")
        path = file_record.get("path", "")
        stem = Path(name).stem
        lesson_id = _filename_lesson_id(stem)

        if not lesson_id:
            report["skipped_files"].append({
                "name": name,
                "lesson_id": "",
                "reason": "filename does not start with a numeric lesson ID",
            })
            continue

        is_ptx = _is_ptx_filename(stem)
        is_transcript = _is_transcript_filename(stem)
        if not is_ptx and not is_transcript:
            report["skipped_files"].append({
                "name": name,
                "lesson_id": lesson_id,
                "reason": "not a *_PTx_NS.docx or *_Transcript.docx file",
            })
            continue

        if lesson_id not in lessons:
            lessons[lesson_id] = {
                "lesson_id": lesson_id,
                "pt_text": "",
                "transcript_text": "",
                "chapter_name": "",
                "pt_filename": "",
                "transcript_filename": "",
            }

        text, read_error = read_docx(path, return_error=True)
        if read_error:
            report["warnings"].append(f"Lesson {lesson_id}: {name} could not be read ({read_error}).")
        elif not text:
            report["warnings"].append(f"Lesson {lesson_id}: {name} has no readable text.")

        if is_ptx:
            if lessons[lesson_id]["pt_text"]:
                lessons[lesson_id]["pt_text"] += "\n\n" + text
                report["warnings"].append(f"Lesson {lesson_id}: duplicate PTx file merged: {name}")
            else:
                lessons[lesson_id]["pt_text"] = text
            lessons[lesson_id]["pt_filename"] = name

            chapter_name = _extract_chapter_name_from_pt(text)
            if chapter_name:
                lessons[lesson_id]["chapter_name"] = chapter_name
                if chapter_name not in report["chapter_names"]:
                    report["chapter_names"].append(chapter_name)
            else:
                report["warnings"].append(
                    f"Lesson {lesson_id}: chapter name was not found in PTx content."
                )

        elif is_transcript:
            if lessons[lesson_id]["transcript_text"]:
                lessons[lesson_id]["transcript_text"] += "\n\n" + text
                report["warnings"].append(f"Lesson {lesson_id}: duplicate Transcript file merged: {name}")
            else:
                lessons[lesson_id]["transcript_text"] = text
            lessons[lesson_id]["transcript_filename"] = name

    if report["chapter_names"]:
        report["chapter_name"] = report["chapter_names"][0]
        if len(report["chapter_names"]) > 1:
            report["warnings"].append(
                "Multiple chapter names found in PTx files: "
                + "; ".join(report["chapter_names"])
            )
        for lesson in lessons.values():
            if not lesson["chapter_name"]:
                # This is not a guessed value; it is the exact chapter name found in another PTx
                # from the same uploaded chapter batch.
                lesson["chapter_name"] = report["chapter_name"]
    else:
        report["warnings"].append(
            "No chapter name was extracted from any PTx file. Generation should not proceed."
        )

    sorted_lessons = sorted(lessons.values(), key=lambda item: int(item["lesson_id"]))
    return (sorted_lessons, report) if return_report else sorted_lessons


# -----------------------------------------------------------------------------
# Budgets and fallback
# -----------------------------------------------------------------------------


def compute_type_budget(num_lessons: int = 0) -> dict[str, int]:
    """7 types x 5 minimum = 35. The remaining 5 slots go to the first 5 types."""
    budget = {question_type: MIN_PER_TYPE for question_type in QUESTION_TYPES}
    flex = QUESTIONS_PER_SET - sum(budget.values())
    for idx in range(flex):
        budget[QUESTION_TYPES[idx]] += 1
    return budget


def compute_fallback_budget(feasible_types: list[str]) -> tuple[dict[str, int], list[str]]:
    """Redistribute 40 slots across feasible types and list dropped types."""
    if "Single Correct" not in feasible_types:
        feasible_types = ["Single Correct"] + [t for t in feasible_types if t != "Single Correct"]
    feasible_types = [t for t in QUESTION_TYPES if t in feasible_types]
    if not feasible_types:
        feasible_types = ["Single Correct"]

    dropped = [t for t in QUESTION_TYPES if t not in feasible_types]
    per_type = QUESTIONS_PER_SET // len(feasible_types)
    remainder = QUESTIONS_PER_SET % len(feasible_types)
    budget = {t: per_type for t in feasible_types}
    for idx in range(remainder):
        budget[feasible_types[idx]] += 1
    return budget, dropped


def _next_fallback_type(short_types: list[str], active_types: list[str]) -> str:
    for question_type in FALLBACK_PRIORITY:
        if question_type in short_types and question_type in active_types:
            return question_type
    for question_type in FALLBACK_PRIORITY:
        if question_type in active_types:
            return question_type
    return ""


# -----------------------------------------------------------------------------
# Prompt construction
# -----------------------------------------------------------------------------

SYSTEM_PROMPT = """You are an expert educational content creator for JoVE.

Absolute rules:
1. Every question must be fully grounded in the supplied source text.
2. Do not invent facts, examples, numbers, terms, definitions, or relationships not present in the source.
3. Do not use the pipe character in question_content.
4. Use the pipe character only where the quiz structure requires it for Dropdown, Match the following, and Categorisation option fields.
5. Return only a valid JSON array. No markdown. No code fences. No preamble.

Each JSON object must have exactly these keys:
lesson_id, question_index, question_content, question_type, option_1, option_2, option_3, option_4, right_answer.

Type rules:

Single Correct:
- Four distinct non-empty options.
- right_answer is one of "1", "2", "3", "4".
- Vary correct-answer position across questions.

Multi Correct:
- Four distinct non-empty options.
- right_answer is comma-separated option numbers.
- Exactly 2 or 3 correct answers. Never 1. Never 4.

True or False:
- option_1 is exactly TRUE.
- option_2 is exactly FALSE.
- option_3 and option_4 are empty strings.
- right_answer is "1" for true or "2" for false.

Fill in the Blanks:
- question_content contains exactly one ___[1]___ placeholder.
- No other blank placeholders.
- option_1 through option_4 are empty strings.
- right_answer is the correct word or phrase.

Dropdown:
- question_content contains 1 or 2 placeholders: ---[dropdown 1]--- and optionally ---[dropdown 2]---.
- option_1 contains exactly 4 pipe-separated choices for dropdown 1.
- option_2 contains exactly 4 pipe-separated choices for dropdown 2 only if there are 2 dropdowns.
- The correct answer is always the first pipe-separated item in each dropdown option field.
- option_3 and option_4 are empty strings.
- right_answer is "1" for one dropdown or "1,2" for two dropdowns.

Match the following:
- Use exactly 3 or 4 pairs.
- Each used option field is exactly: Left term | Right term.
- Each used option field has exactly one pipe character.
- right_answer is an empty string.

Categorisation:
- Use 2 to 4 categories.
- Each used option field is: Category Name | item1 | item2, with 2 to 4 items.
- Each category option therefore has 2 to 4 pipe-separated items after the category name.
- right_answer is an empty string.
"""


def _source_for_prompt(lesson: dict[str, Any]) -> str:
    parts: list[str] = []
    if lesson.get("transcript_text"):
        parts.append("[TRANSCRIPT - PRIMARY SOURCE]\n" + lesson["transcript_text"])
    if lesson.get("pt_text"):
        parts.append("[PTx - SUPPORTING SOURCE]\n" + lesson["pt_text"])
    return "\n\n".join(parts)


def _format_budget_instruction(budget: dict[str, int]) -> str:
    total = sum(budget.values())
    lines = [f'- {count} questions of type "{question_type}"' for question_type, count in budget.items()]
    return "Generate exactly " + str(total) + " questions with this exact breakdown:\n" + "\n".join(lines)


def build_set_prompt(
    lessons: list[dict[str, Any]],
    set_number: int,
    budget: dict[str, int],
    previously_used_snippets: list[str],
    start_index: int = 1,
) -> str:
    source_blocks = []
    for lesson in lessons:
        source = _source_for_prompt(lesson)
        if source:
            source_blocks.append(f"LESSON {lesson['lesson_id']}\n{source}")

    if previously_used_snippets:
        exclusion = (
            "Do not duplicate or closely paraphrase these already-used questions:\n"
            + "\n".join(f"- {snippet}" for snippet in previously_used_snippets[:200])
        )
    else:
        exclusion = "This is Set 1. There are no previously used questions."

    lesson_ids = ", ".join(lesson["lesson_id"] for lesson in lessons)
    return f"""Generate quiz Set {set_number} of {NUM_SETS} for chapter: "{lessons[0]['chapter_name']}".

Allowed lesson_id values: {lesson_ids}
Use only the source material below. Transcript is primary when available. If a lesson has no transcript, use its PTx as the source.

SOURCE MATERIAL:
{chr(10).join(source_blocks)}

EXCLUSION CONTEXT:
{exclusion}

{_format_budget_instruction(budget)}

Additional requirements:
- question_index starts at {start_index} and increments by 1.
- Spread questions across the allowed lessons; do not let one lesson dominate the set.
- Each generated object must use a lesson_id from the allowed list.
- Each question must test a distinct concept.
- Return only the JSON array.
"""


def build_topup_prompt(
    lessons: list[dict[str, Any]],
    set_number: int,
    needed: dict[str, int],
    existing_snippets: list[str],
    next_index: int,
) -> str:
    source_blocks = []
    for lesson in lessons:
        source = _source_for_prompt(lesson)
        if source:
            source_blocks.append(f"LESSON {lesson['lesson_id']}\n{source}")

    lesson_ids = ", ".join(lesson["lesson_id"] for lesson in lessons)
    existing = "\n".join(f"- {snippet}" for snippet in existing_snippets[:240])
    return f"""Top-up generation for Set {set_number}.

Allowed lesson_id values: {lesson_ids}
Generate only the missing question types and counts listed below. Do not generate any other type.

SOURCE MATERIAL:
{chr(10).join(source_blocks)}

ALREADY GENERATED OR REJECTED FOR DUPLICATION:
{existing}

{_format_budget_instruction(needed)}

Additional requirements:
- question_index starts at {next_index}.
- Follow every type rule exactly.
- Do not duplicate or closely paraphrase existing questions.
- Return only the JSON array.
"""


# -----------------------------------------------------------------------------
# LLM call and JSON parsing
# -----------------------------------------------------------------------------


class LLMRequestError(RuntimeError):
    """Raised when the OpenAI request fails before any usable model output is returned."""


def _looks_like_unsupported_param_error(exc: Exception, param_name: str) -> bool:
    message = str(exc).lower()
    return param_name.lower() in message and (
        "unsupported" in message or "unrecognized" in message or "not supported" in message
    )


def call_llm(system: str, user: str, api_key: str, model: str = "gpt-4o") -> str:
    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]
    request_args = {
        "model": model,
        "messages": messages,
        "temperature": 0.65,
    }
    completion_limit = 16000

    try:
        from openai import OpenAI  # type: ignore
    except (ImportError, AttributeError) as exc:
        raise LLMRequestError(
            "OpenAI Python SDK is missing or incompatible. Install/upgrade the openai package."
        ) from exc

    client = OpenAI(api_key=api_key)
    try:
        try:
            response = client.chat.completions.create(
                **request_args,
                max_completion_tokens=completion_limit,
            )
        except Exception as exc:
            # Prefer max_completion_tokens because newer reasoning models reject max_tokens.
            # Fall back only when max_completion_tokens itself is rejected by an older/non-standard model.
            if not _looks_like_unsupported_param_error(exc, "max_completion_tokens"):
                raise
            response = client.chat.completions.create(
                **request_args,
                max_tokens=completion_limit,
            )
    except Exception as exc:
        raise LLMRequestError(str(exc)) from exc

    return response.choices[0].message.content or ""


def parse_llm_json(raw: str) -> list[dict[str, Any]]:
    """Strip common fences and parse a JSON array."""
    cleaned = re.sub(r"```(?:json)?", "", raw).strip().strip("`").strip()
    start = cleaned.find("[")
    end = cleaned.rfind("]")
    if start == -1 or end == -1 or end <= start:
        raise ValueError(f"No JSON array found in response. First 300 chars: {raw[:300]}")
    parsed = json.loads(cleaned[start:end + 1])
    if not isinstance(parsed, list):
        raise ValueError("LLM response JSON was not an array.")
    return parsed


# -----------------------------------------------------------------------------
# Validation helpers
# -----------------------------------------------------------------------------


def _as_string(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _norm_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _norm_for_compare(text: str) -> str:
    text = text.lower()
    text = re.sub(r"---\[dropdown\s*\d+\]---", " dropdownplaceholder ", text)
    text = re.sub(r"___\[1\]___", " blankplaceholder ", text)
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return _norm_space(text)


def _hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _full_question_text(q: dict[str, Any]) -> str:
    return " ".join(
        _as_string(q.get(field, ""))
        for field in ["question_content", "option_1", "option_2", "option_3", "option_4"]
    )


def _question_content_signature(q: dict[str, Any]) -> str:
    return _hash_text(_norm_for_compare(_as_string(q.get("question_content", ""))))


def _full_question_signature(q: dict[str, Any]) -> str:
    return _hash_text(_norm_for_compare(_full_question_text(q)))


def _question_norm(q: dict[str, Any]) -> str:
    return _norm_for_compare(_as_string(q.get("question_content", "")))


def _snippet(q: dict[str, Any], max_chars: int = EXCLUSION_SNIPPET_CHARS) -> str:
    text = _norm_space(_full_question_text(q))
    if len(text) > max_chars:
        text = text[:max_chars].rstrip() + "..."
    return f"[{q.get('question_type', '?')}] {text}"


def _split_pipe(value: str) -> list[str]:
    return [part.strip() for part in value.split("|")]


def _parse_multi_answer(value: str) -> list[str]:
    parts = [part.strip() for part in value.split(",") if part.strip()]
    return parts


def _all_distinct(options: list[str]) -> bool:
    normalized = [_norm_for_compare(option) for option in options]
    return len(normalized) == len(set(normalized))


def _significant_terms(text: str) -> set[str]:
    normalized = re.sub(r"[^A-Za-z0-9]+", " ", text).lower()
    terms = set()
    for token in normalized.split():
        if len(token) >= 4 and token not in STOPWORDS and not token.isdigit():
            terms.add(token)
    return terms


def _lesson_source_text(lesson: dict[str, Any]) -> str:
    # Transcript is primary, PTx is supporting. If transcript is empty, PTx is the sole source.
    pieces = []
    if lesson.get("transcript_text"):
        pieces.append(lesson["transcript_text"])
    if lesson.get("pt_text"):
        pieces.append(lesson["pt_text"])
    return "\n".join(pieces)


def _grounding_check(q: dict[str, Any], lesson: dict[str, Any]) -> tuple[list[str], list[str]]:
    """
    Lightweight deterministic grounding check.

    This does not prove semantic truth. It catches obvious source drift by checking that
    the important terms in the question/options overlap with the selected lesson source.
    """
    warnings: list[str] = []
    errors: list[str] = []
    source_text = _lesson_source_text(lesson)
    source_terms = _significant_terms(source_text)
    question_terms = _significant_terms(_full_question_text(q))

    if not source_terms:
        errors.append("selected lesson has no source text for grounding validation")
        return warnings, errors
    if not question_terms:
        warnings.append("no significant terms found for grounding validation")
        return warnings, errors

    overlap = question_terms.intersection(source_terms)
    ratio = len(overlap) / max(1, len(question_terms))
    missing = sorted(question_terms.difference(source_terms))[:12]
    if ratio < GROUNDING_MIN_OVERLAP:
        errors.append(
            "low source-term overlap for grounding validation; missing terms include: "
            + ", ".join(missing)
        )
    elif missing:
        warnings.append(
            "grounding check flagged terms not found verbatim in selected lesson: "
            + ", ".join(missing)
        )
    return warnings, errors


def validate_and_repair(
    q: dict[str, Any],
    index: int,
    allowed_lesson_ids: set[str] | None = None,
    lessons_by_id: dict[str, dict[str, Any]] | None = None,
) -> tuple[dict[str, Any], list[str], list[str]]:
    """
    Validate and minimally repair one question.

    Returns (clean_question, warnings, errors). Any errors mean the question must not be exported.
    """
    warnings: list[str] = []
    errors: list[str] = []

    if not isinstance(q, dict):
        return {}, [], [f"Q{index}: item is not a JSON object"]

    clean: dict[str, Any] = {field: _as_string(q.get(field, "")) for field in REQUIRED_FIELDS}
    clean["question_index"] = index

    lesson_id = clean["lesson_id"]
    if allowed_lesson_ids is not None and lesson_id not in allowed_lesson_ids:
        if len(allowed_lesson_ids) == 1:
            repaired_id = next(iter(allowed_lesson_ids))
            warnings.append(f"Q{index}: invalid or blank lesson_id '{lesson_id}' repaired to {repaired_id}")
            lesson_id = repaired_id
            clean["lesson_id"] = repaired_id
        else:
            errors.append(f"Q{index}: lesson_id '{lesson_id}' is not in uploaded lesson IDs")

    question_type = clean["question_type"]
    question_content = clean["question_content"]

    if question_type not in QUESTION_TYPES:
        errors.append(f"Q{index}: unknown question_type '{question_type}'")
    if not question_content:
        errors.append(f"Q{index} [{question_type}]: question_content is empty")
    if "|" in question_content:
        errors.append(f"Q{index} [{question_type}]: question_content contains illegal pipe character")

    for option_key in ["option_1", "option_2", "option_3", "option_4"]:
        if clean[option_key].lower() == "none":
            clean[option_key] = ""
            warnings.append(f"Q{index} [{question_type}]: literal None removed from {option_key}")

    if question_type in {"Single Correct", "Multi Correct", "True or False", "Fill in the Blanks"}:
        for option_key in ["option_1", "option_2", "option_3", "option_4"]:
            if "|" in clean[option_key]:
                errors.append(f"Q{index} [{question_type}]: {option_key} contains illegal pipe character")

    if question_type == "Single Correct":
        options = [clean[f"option_{idx}"] for idx in range(1, 5)]
        if any(not option for option in options):
            errors.append(f"Q{index} [Single Correct]: exactly 4 non-empty options are required")
        if not _all_distinct(options):
            errors.append(f"Q{index} [Single Correct]: options must be distinct")
        if clean["right_answer"] not in {"1", "2", "3", "4"}:
            errors.append(f"Q{index} [Single Correct]: right_answer must be 1, 2, 3, or 4")

    elif question_type == "Multi Correct":
        options = [clean[f"option_{idx}"] for idx in range(1, 5)]
        if any(not option for option in options):
            errors.append(f"Q{index} [Multi Correct]: exactly 4 non-empty options are required")
        if not _all_distinct(options):
            errors.append(f"Q{index} [Multi Correct]: options must be distinct")
        answers = _parse_multi_answer(clean["right_answer"])
        if len(answers) != len(set(answers)):
            errors.append(f"Q{index} [Multi Correct]: duplicate answer numbers are not allowed")
        if any(answer not in {"1", "2", "3", "4"} for answer in answers):
            errors.append(f"Q{index} [Multi Correct]: right_answer contains invalid option numbers")
        if len(answers) not in {2, 3}:
            errors.append(f"Q{index} [Multi Correct]: right_answer must contain exactly 2 or 3 options")
        if not errors:
            clean["right_answer"] = ",".join(sorted(answers, key=int))

    elif question_type == "True or False":
        clean["option_1"] = "TRUE"
        clean["option_2"] = "FALSE"
        clean["option_3"] = ""
        clean["option_4"] = ""
        if clean["right_answer"] not in {"1", "2"}:
            errors.append(f"Q{index} [True or False]: right_answer must be 1 or 2")

    elif question_type == "Fill in the Blanks":
        clean["option_1"] = ""
        clean["option_2"] = ""
        clean["option_3"] = ""
        clean["option_4"] = ""
        exact_count = question_content.count("___[1]___")
        all_placeholders = re.findall(r"___\[\d+\]___", question_content)
        if exact_count != 1 or len(all_placeholders) != 1:
            errors.append(
                f"Q{index} [Fill in the Blanks]: question_content must contain exactly one ___[1]___ placeholder"
            )
        if not clean["right_answer"]:
            errors.append(f"Q{index} [Fill in the Blanks]: right_answer is required")

    elif question_type == "Dropdown":
        placeholder_matches = list(re.finditer(r"---\[dropdown\s+(\d+)\]---", question_content))
        placeholder_numbers = [int(match.group(1)) for match in placeholder_matches]
        if placeholder_numbers not in ([1], [1, 2]):
            errors.append(
                f"Q{index} [Dropdown]: dropdown placeholders must be numbered exactly [1] or [1, 2]; got {placeholder_numbers or 'none'}"
            )
        n_dropdowns = len(placeholder_numbers)
        if n_dropdowns not in {1, 2}:
            errors.append(f"Q{index} [Dropdown]: question_content must contain 1 or 2 dropdown placeholders")
        if n_dropdowns >= 1:
            expected_answer = "1" if n_dropdowns == 1 else "1,2"
            if clean["right_answer"] != expected_answer:
                warnings.append(
                    f"Q{index} [Dropdown]: right_answer repaired from '{clean['right_answer']}' to '{expected_answer}'"
                )
                clean["right_answer"] = expected_answer
        for idx in range(1, min(n_dropdowns, 2) + 1):
            option_key = f"option_{idx}"
            parts = _split_pipe(clean[option_key])
            if len(parts) != 4 or any(not part for part in parts):
                errors.append(f"Q{index} [Dropdown]: {option_key} must have exactly 4 non-empty pipe-separated choices")
            else:
                clean[option_key] = " | ".join(parts)
        for idx in range(min(n_dropdowns, 2) + 1, 5):
            option_key = f"option_{idx}"
            if clean[option_key]:
                errors.append(f"Q{index} [Dropdown]: {option_key} must be empty")

    elif question_type == "Match the following":
        clean["right_answer"] = ""
        used_options = [clean[f"option_{idx}"] for idx in range(1, 5) if clean[f"option_{idx}"]]
        if len(used_options) not in {3, 4}:
            errors.append(f"Q{index} [Match the following]: exactly 3 or 4 pairs are required")
        for option_number, option_value in enumerate([clean[f"option_{idx}"] for idx in range(1, 5)], 1):
            if not option_value:
                continue
            parts = _split_pipe(option_value)
            if len(parts) != 2 or any(not part for part in parts):
                errors.append(
                    f"Q{index} [Match the following]: option_{option_number} must be 'Left term | Right term' with exactly one pipe"
                )
            else:
                clean[f"option_{option_number}"] = " | ".join(parts)

    elif question_type == "Categorisation":
        clean["right_answer"] = ""
        used_options = [clean[f"option_{idx}"] for idx in range(1, 5) if clean[f"option_{idx}"]]
        if len(used_options) < 2 or len(used_options) > 4:
            errors.append(f"Q{index} [Categorisation]: 2 to 4 categories are required")
        for option_number, option_value in enumerate([clean[f"option_{idx}"] for idx in range(1, 5)], 1):
            if not option_value:
                continue
            parts = _split_pipe(option_value)
            item_count = len(parts) - 1
            if len(parts) < 3 or len(parts) > 5 or any(not part for part in parts):
                errors.append(
                    f"Q{index} [Categorisation]: option_{option_number} must be category plus 2 to 4 non-empty items"
                )
            elif item_count < 2 or item_count > 4:
                errors.append(
                    f"Q{index} [Categorisation]: option_{option_number} must contain 2 to 4 items"
                )
            else:
                clean[f"option_{option_number}"] = " | ".join(parts)

    if lessons_by_id is not None and lesson_id in lessons_by_id:
        ground_warnings, ground_errors = _grounding_check(clean, lessons_by_id[lesson_id])
        warnings.extend(f"Q{index} [{question_type}]: {msg}" for msg in ground_warnings)
        errors.extend(f"Q{index} [{question_type}]: {msg}" for msg in ground_errors)

    return clean, warnings, errors


# -----------------------------------------------------------------------------
# Acceptance, counting, duplicate prevention, and balancing
# -----------------------------------------------------------------------------


def _count_by_type(questions: list[dict[str, Any]]) -> dict[str, int]:
    counts = {question_type: 0 for question_type in QUESTION_TYPES}
    for question in questions:
        question_type = question.get("question_type", "")
        if question_type in counts:
            counts[question_type] += 1
    return counts


def _count_by_lesson(questions: list[dict[str, Any]], lesson_ids: list[str]) -> dict[str, int]:
    counts = {lesson_id: 0 for lesson_id in lesson_ids}
    for question in questions:
        lesson_id = str(question.get("lesson_id", ""))
        if lesson_id in counts:
            counts[lesson_id] += 1
    return counts


def _token_jaccard(left: str, right: str) -> float:
    left_tokens = set(left.split())
    right_tokens = set(right.split())
    if not left_tokens or not right_tokens:
        return 0.0
    return len(left_tokens.intersection(right_tokens)) / len(left_tokens.union(right_tokens))


def _is_near_duplicate(candidate_norm: str, existing_norms: list[str]) -> bool:
    if not candidate_norm or len(candidate_norm) < SHORT_NEAR_DUPLICATE_MIN_CHARS:
        return False
    for existing in existing_norms:
        if not existing:
            continue
        ratio = difflib.SequenceMatcher(None, candidate_norm, existing).ratio()
        jaccard = _token_jaccard(candidate_norm, existing)
        if len(candidate_norm) < 24 or len(existing) < 24:
            # Short questions often differ by only a plural, tense, or omitted word;
            # token Jaccard is too brittle at this length, so rely mainly on
            # character similarity while exact hashes still catch verbatim copies.
            if ratio >= SHORT_NEAR_DUPLICATE_THRESHOLD:
                return True
        elif ratio >= NEAR_DUPLICATE_THRESHOLD and jaccard >= 0.88:
            return True
    return False


def _dedup_reason(
    candidate: dict[str, Any],
    used_content_signatures: set[str],
    used_full_signatures: set[str],
    used_norms: list[str],
) -> str:
    content_sig = _question_content_signature(candidate)
    full_sig = _full_question_signature(candidate)
    candidate_norm = _question_norm(candidate)
    if content_sig in used_content_signatures:
        return "duplicate question_content"
    if full_sig in used_full_signatures:
        return "duplicate full question/options"
    if _is_near_duplicate(candidate_norm, used_norms):
        return "near-duplicate question_content"
    return ""


def _register_question(
    question: dict[str, Any],
    used_content_signatures: set[str],
    used_full_signatures: set[str],
    used_norms: list[str],
) -> None:
    used_content_signatures.add(_question_content_signature(question))
    used_full_signatures.add(_full_question_signature(question))
    used_norms.append(_question_norm(question))


def _build_dedup_state(questions: list[dict[str, Any]]) -> tuple[set[str], set[str], list[str]]:
    content_signatures: set[str] = set()
    full_signatures: set[str] = set()
    norms: list[str] = []
    for question in questions:
        _register_question(question, content_signatures, full_signatures, norms)
    return content_signatures, full_signatures, norms


def _reindex(questions: list[dict[str, Any]]) -> None:
    for idx, question in enumerate(questions, 1):
        question["question_index"] = idx


def _needed_counts(questions: list[dict[str, Any]], budget: dict[str, int]) -> dict[str, int]:
    counts = _count_by_type(questions)
    needed = {}
    for question_type, required in budget.items():
        have = counts.get(question_type, 0)
        if have < required:
            needed[question_type] = required - have
    return needed


def _accept_candidates(
    candidates: list[dict[str, Any]],
    set_questions: list[dict[str, Any]],
    budget: dict[str, int],
    wanted_types: set[str] | None,
    allowed_lesson_ids: set[str],
    lessons_by_id: dict[str, dict[str, Any]],
    used_content_signatures: set[str],
    used_full_signatures: set[str],
    used_norms: list[str],
    set_warnings: list[str],
    phase: str,
) -> int:
    accepted = 0
    for raw_index, raw_question in enumerate(candidates, 1):
        if len(set_questions) >= QUESTIONS_PER_SET:
            set_warnings.append(f"{phase}: extra question rejected because set already has 40 accepted questions")
            continue

        candidate_index = len(set_questions) + 1
        clean, warnings, errors = validate_and_repair(
            raw_question,
            candidate_index,
            allowed_lesson_ids=allowed_lesson_ids,
            lessons_by_id=lessons_by_id,
        )
        set_warnings.extend(warnings)
        if errors:
            set_warnings.extend(f"{phase}: rejected - {error}" for error in errors)
            continue

        question_type = clean["question_type"]
        if wanted_types is not None and question_type not in wanted_types:
            set_warnings.append(
                f"{phase}: rejected Q{raw_index}; returned type '{question_type}' was not requested in this top-up"
            )
            continue

        type_counts = _count_by_type(set_questions)
        if type_counts.get(question_type, 0) >= budget.get(question_type, 0):
            set_warnings.append(
                f"{phase}: rejected Q{raw_index}; type '{question_type}' is already at quota"
            )
            continue

        duplicate_reason = _dedup_reason(
            clean,
            used_content_signatures,
            used_full_signatures,
            used_norms,
        )
        if duplicate_reason:
            set_warnings.append(f"{phase}: rejected Q{raw_index}; {duplicate_reason}")
            continue

        set_questions.append(clean)
        _register_question(clean, used_content_signatures, used_full_signatures, used_norms)
        accepted += 1

    return accepted


def _lesson_distribution_warnings(
    questions: list[dict[str, Any]],
    lesson_ids: list[str],
    set_number: int,
) -> list[str]:
    if len(lesson_ids) <= 1 or not questions:
        return []
    warnings: list[str] = []
    counts = _count_by_lesson(questions, lesson_ids)
    expected_per_lesson = len(questions) / len(lesson_ids)
    buffer = max(1, int(expected_per_lesson * 0.05))
    max_allowed = int(expected_per_lesson + buffer)
    for lesson_id, count in counts.items():
        if count == 0:
            warnings.append(f"Set {set_number}: Lesson {lesson_id} has no accepted questions.")
        elif count > max_allowed:
            warnings.append(
                f"Set {set_number}: Lesson {lesson_id} dominates the set with {count} questions; expected about {expected_per_lesson:.1f}, allowed maximum {max_allowed}."
            )
    return warnings


def _single_correct_position_warnings(questions: list[dict[str, Any]], set_number: int) -> list[str]:
    sc_answers = [q.get("right_answer", "") for q in questions if q.get("question_type") == "Single Correct"]
    if len(sc_answers) < 2:
        return []
    counts = {answer: sc_answers.count(answer) for answer in {"1", "2", "3", "4"}}
    dominant_answer, dominant_count = max(counts.items(), key=lambda item: item[1])
    dominance_ratio = dominant_count / len(sc_answers)
    if len(set(sc_answers)) == 1:
        return [
            f"Set {set_number}: all Single Correct answers use option {dominant_answer}; answer position is not varied."
        ]
    if dominance_ratio > 0.60:
        return [
            f"Set {set_number}: Single Correct answer option {dominant_answer} dominates ({dominant_count}/{len(sc_answers)} = {dominance_ratio:.0%}); answer positions should be more varied."
        ]
    return []


def _notify(callback: Callable[..., None] | None, message: str, progress: int | None = None) -> None:
    if not callback:
        return
    try:
        callback(message, progress)
    except TypeError:
        callback(message)


# -----------------------------------------------------------------------------
# Generation pipeline
# -----------------------------------------------------------------------------


def generate_quiz(
    lessons: list[dict[str, Any]],
    api_key: str,
    model: str = "gpt-4o",
    progress_callback: Callable[..., None] | None = None,
    all_lessons: list[dict[str, Any]] | None = None,
    skipped_lessons: list[str] | None = None,
    parse_warnings: list[str] | None = None,
    skipped_files: list[dict[str, str]] | None = None,
    chapter_name: str | None = None,
) -> tuple[list[list[dict[str, Any]]], dict[str, Any]]:
    """Generate three quiz sets and return (all_sets, report)."""
    if not lessons:
        raise ValueError("No content-bearing lessons were supplied.")

    allowed_lesson_ids = {lesson["lesson_id"] for lesson in lessons}
    lessons_by_id = {lesson["lesson_id"]: lesson for lesson in lessons}
    ordered_lesson_ids = [lesson["lesson_id"] for lesson in lessons]
    resolved_chapter_name = chapter_name or lessons[0].get("chapter_name", "")
    if not resolved_chapter_name:
        raise ValueError("Chapter name was not extracted from PTx content.")

    all_sets: list[list[dict[str, Any]]] = []
    all_warnings: list[str] = list(parse_warnings or [])
    report: dict[str, Any] = {
        "chapter_name": resolved_chapter_name,
        "num_lessons": len(all_lessons) if all_lessons is not None else len(lessons),
        "num_content_lessons": len(lessons),
        "skipped_lessons": skipped_lessons or [],
        "skipped_files": skipped_files or [],
        "sets": [],
        "warnings": [],
        "fallback_types_dropped": [],
    }

    used_content_signatures: set[str] = set()
    used_full_signatures: set[str] = set()
    used_norms: list[str] = []
    all_generated: list[dict[str, Any]] = []
    globally_dropped_types: set[str] = set()

    base_budget = compute_type_budget(len(lessons))

    for set_number in range(1, NUM_SETS + 1):
        set_questions: list[dict[str, Any]] = []
        set_warnings: list[str] = []
        active_types = list(QUESTION_TYPES)
        current_budget = dict(base_budget)
        set_dropped_types: list[str] = []

        base_progress = 5 + int((set_number - 1) * 25)
        _notify(progress_callback, f"Generating Set {set_number}/{NUM_SETS}: initial pass", base_progress)

        try:
            initial_prompt = build_set_prompt(
                lessons=lessons,
                set_number=set_number,
                budget=current_budget,
                previously_used_snippets=[_snippet(q) for q in all_generated],
                start_index=1,
            )
            raw = call_llm(SYSTEM_PROMPT, initial_prompt, api_key, model)
            raw_questions = parse_llm_json(raw)
        except LLMRequestError as exc:
            raise RuntimeError(f"OpenAI request failed before question generation: {exc}") from exc
        except Exception as exc:
            set_warnings.append(f"Set {set_number}: initial generation failed or returned invalid JSON: {exc}")
            raw_questions = []

        _accept_candidates(
            candidates=raw_questions,
            set_questions=set_questions,
            budget=current_budget,
            wanted_types=None,
            allowed_lesson_ids=allowed_lesson_ids,
            lessons_by_id=lessons_by_id,
            used_content_signatures=used_content_signatures,
            used_full_signatures=used_full_signatures,
            used_norms=used_norms,
            set_warnings=set_warnings,
            phase=f"Set {set_number} initial pass",
        )

        # Top-up and fallback. The outer loop allows fallback to recompute quotas.
        fallback_rounds = 0
        while True:
            needed = _needed_counts(set_questions, current_budget)
            if not needed and len(set_questions) == QUESTIONS_PER_SET:
                break
            if not needed and len(set_questions) < QUESTIONS_PER_SET:
                set_warnings.append(
                    f"Set {set_number}: internal budget inconsistency before top-up; quotas are satisfied but only {len(set_questions)} questions were accepted."
                )
                break

            for attempt in range(1, MAX_RETRIES + 1):
                needed = _needed_counts(set_questions, current_budget)
                if not needed and len(set_questions) == QUESTIONS_PER_SET:
                    break
                if not needed and len(set_questions) < QUESTIONS_PER_SET:
                    set_warnings.append(
                        f"Set {set_number}: internal budget inconsistency; quotas are satisfied but only {len(set_questions)} questions were accepted. No defensive budget mutation was applied."
                    )
                    break

                shortage = ", ".join(f"{question_type}: {count}" for question_type, count in needed.items())
                _notify(
                    progress_callback,
                    f"Set {set_number} top-up attempt {attempt}/{MAX_RETRIES}: {shortage}",
                    min(84, base_progress + attempt * 5 + fallback_rounds * 3),
                )

                try:
                    topup_prompt = build_topup_prompt(
                        lessons=lessons,
                        set_number=set_number,
                        needed=needed,
                        existing_snippets=[_snippet(q) for q in all_generated + set_questions],
                        next_index=len(set_questions) + 1,
                    )
                    raw_topup = call_llm(SYSTEM_PROMPT, topup_prompt, api_key, model)
                    topup_questions = parse_llm_json(raw_topup)
                except LLMRequestError as exc:
                    raise RuntimeError(f"OpenAI request failed during top-up attempt {attempt}: {exc}") from exc
                except Exception as exc:
                    set_warnings.append(f"Set {set_number}: top-up attempt {attempt} failed: {exc}")
                    topup_questions = []

                _accept_candidates(
                    candidates=topup_questions,
                    set_questions=set_questions,
                    budget=current_budget,
                    wanted_types=set(needed.keys()),
                    allowed_lesson_ids=allowed_lesson_ids,
                    lessons_by_id=lessons_by_id,
                    used_content_signatures=used_content_signatures,
                    used_full_signatures=used_full_signatures,
                    used_norms=used_norms,
                    set_warnings=set_warnings,
                    phase=f"Set {set_number} top-up attempt {attempt}",
                )

            needed_after_retries = _needed_counts(set_questions, current_budget)
            if not needed_after_retries and len(set_questions) == QUESTIONS_PER_SET:
                break

            short_types = list(needed_after_retries.keys())
            fallback_type = _next_fallback_type(short_types, active_types)
            if not fallback_type:
                break

            fallback_rounds += 1
            active_types.remove(fallback_type)
            set_dropped_types.append(fallback_type)
            globally_dropped_types.add(fallback_type)
            set_warnings.append(
                f"Set {set_number}: fallback applied; dropped '{fallback_type}' after {MAX_RETRIES} top-up attempts."
            )

            # Remove already accepted questions of the dropped type so the Summary accurately
            # reports it as not generated for the fallback set.
            retained_questions = [q for q in set_questions if q.get("question_type") != fallback_type]
            dropped_count = len(set_questions) - len(retained_questions)
            if dropped_count:
                set_warnings.append(
                    f"Set {set_number}: removed {dropped_count} previously accepted '{fallback_type}' questions after fallback."
                )
            set_questions = retained_questions
            used_content_signatures, used_full_signatures, used_norms = _build_dedup_state(all_generated + set_questions)

            current_budget, dropped_by_budget = compute_fallback_budget(active_types)
            for dropped in dropped_by_budget:
                globally_dropped_types.add(dropped)
            _reindex(set_questions)

            if len(active_types) == 1 and active_types[0] == "Single Correct" and fallback_rounds > len(FALLBACK_PRIORITY):
                break

            if fallback_rounds > len(FALLBACK_PRIORITY):
                break

        final_counts = _count_by_type(set_questions)
        if len(set_questions) != QUESTIONS_PER_SET:
            set_warnings.append(
                f"Set {set_number}: generated {len(set_questions)} accepted questions; expected exactly {QUESTIONS_PER_SET}."
            )

        for question_type, required in current_budget.items():
            have = final_counts.get(question_type, 0)
            if have < required:
                set_warnings.append(
                    f"Set {set_number}: {question_type} shortfall after top-up/fallback; have {have}, need {required}."
                )
            elif have > required:
                set_warnings.append(
                    f"Set {set_number}: {question_type} over quota; have {have}, expected {required}."
                )

        set_warnings.extend(_lesson_distribution_warnings(set_questions, ordered_lesson_ids, set_number))
        set_warnings.extend(_single_correct_position_warnings(set_questions, set_number))

        _reindex(set_questions)
        all_sets.append(set_questions)
        all_generated.extend(set_questions)
        all_warnings.extend(set_warnings)

        report["sets"].append({
            "set_number": set_number,
            "total_questions": len(set_questions),
            "type_counts": final_counts,
            "budget": current_budget,
            "dropped_types": set_dropped_types,
            "lesson_counts": _count_by_lesson(set_questions, ordered_lesson_ids),
            "warnings": set_warnings,
        })

    report["fallback_types_dropped"] = [
        question_type for question_type in QUESTION_TYPES if question_type in globally_dropped_types
    ]
    report["warnings"] = all_warnings
    return all_sets, report
